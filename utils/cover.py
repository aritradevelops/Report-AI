from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.styles.styles import WD_STYLE_TYPE
from htmldocx import HtmlToDocx
from docx2pdf import convert
new_parser = HtmlToDocx()

def add_bold_label(self, label, value, size=20):
    para = self.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    # Add bold label with size
    run_label = para.add_run(label)
    run_label.bold = True
    run_label.font.size = Pt(size)
    
    # Add normal value with the same size
    para_run = para.add_run(value)
    para_run.font.size = Pt(size)

def generate_doc_with_intro(topic, labels, gap=5):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    # font.name = 'Helvatica'
    font.size = Pt(18)

    # Add title with bigger font size and center alignment
    title = doc.add_paragraph()
    title_run = title.add_run(topic)
    title_run.font.size = Pt(40)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment for title
    
    # Add horizontal rule
    new_parser.add_html_to_document('<hr/>', doc)

    # Add couple of empty lines here
    for _ in range(gap):
        doc.add_paragraph("\n")

    # Bind the add_bold_label method to the Document instance
    doc.add_bold_label = add_bold_label.__get__(doc)

    # Add content with bold labels and normal values
    for label, value in labels.items():
        doc.add_bold_label(label + ': ', value + '.')

    return doc

# Usage example:
if __name__ == '__main__':
    doc = generate_doc_with_intro(
        "Mathematics III",
        {
            "name": "Aritra Sadhukhan",
            "roll_no": "27900123051",
            "dept": "Computer Science & Engineering",
            "subject": "Mathematics III",
            "subject_code": "BSC301",
            "institute": "Ideal Institute Of Engineering"
        }
    )

